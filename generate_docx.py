#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar DOCX de la documentación README.md
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Crear documento
doc = Document()

# Título principal
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run("NOVA API IA")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 119, 180)

# Subtítulo
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run("Documentación Completa del Proyecto")
run.font.size = Pt(14)

# Fecha
fecha = doc.add_paragraph()
fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = fecha.add_run(f"Generado: {datetime.now().strftime('%d de %B de %Y')}")
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# Descripción del Proyecto
doc.add_heading("Descripción del Proyecto", level=1)
doc.add_paragraph(
    "NOVA API IA es una API REST construida con Flask que proporciona acceso a modelos de inteligencia artificial "
    "locales mediante Ollama. Esta API permite que la aplicación NOVA realice consultas a modelos de IA sin depender "
    "de servicios en la nube."
)

doc.add_heading("Características Principales:", level=2)
features = [
    "API REST con endpoints simples y claros",
    "Integración con Ollama (ejecución local de modelos IA)",
    "Modelo Qwen 2.5 0.5B (ligero y rápido)",
    "Sistema completo de logs y monitoreo",
    "CORS habilitado para compatibilidad multi-origen",
    "Health check automático",
    "Manejo robusto de errores"
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# Requisitos del Sistema
doc.add_heading("Requisitos del Sistema", level=1)

doc.add_heading("Software Requerido:", level=2)
requirements = [
    "Python 3.8 o superior",
    "Ollama (para ejecutar modelos de IA)",
    "pip (gestor de paquetes de Python)",
    "Windows PowerShell o CMD"
]
for req in requirements:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading("Requisitos Mínimos de Hardware:", level=2)
hw = [
    "Procesador: Intel i5 / AMD Ryzen 5 o superior",
    "RAM: 4 GB mínimo, 8 GB recomendado",
    "Almacenamiento: 5 GB libres (para modelos)",
    "Conexión a Internet (solo para descargas iniciales)"
]
for h in hw:
    doc.add_paragraph(h, style='List Bullet')

# Instalación y Configuración
doc.add_heading("Instalación y Configuración", level=1)

doc.add_heading("Paso 1: Navegar a la Carpeta del Proyecto", level=2)
p = doc.add_paragraph()
run = p.add_run("cd C:\\xampp\\htdocs\\NOVA_ASISTENTE\\API_IA")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Paso 2: Crear Entorno Virtual", level=2)
p = doc.add_paragraph()
run = p.add_run("python -m venv venv")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Paso 3: Activar Entorno Virtual", level=2)
p = doc.add_paragraph()
run = p.add_run(".\\venv\\Scripts\\activate")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Paso 4: Instalar Dependencias", level=2)
p = doc.add_paragraph()
run = p.add_run("pip install -r requirements.txt")
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Ollama
doc.add_heading("Configuración de Ollama e IA", level=1)

doc.add_heading("Paso 1: Descargar e Instalar Ollama", level=2)
doc.add_paragraph("1. Ir a https://ollama.ai", style='List Number')
doc.add_paragraph("2. Descargar la versión para Windows", style='List Number')
doc.add_paragraph("3. Ejecutar el instalador y seguir las instrucciones", style='List Number')

doc.add_heading("Paso 2: Descargar Modelo Qwen 2.5 0.5B", level=2)
doc.add_paragraph("Ejecutar en PowerShell (sin venv activado):")
p = doc.add_paragraph()
run = p.add_run("ollama pull qwen2.5:0.5b")
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph("Esto descargará aproximadamente 300-400 MB. Tiempo estimado: 2-5 minutos.")

doc.add_heading("Paso 3: Iniciar Ollama en Modo Servidor", level=2)
doc.add_paragraph("En una terminal nueva ejecutar:")
p = doc.add_paragraph()
run = p.add_run("ollama serve")
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph("Mantén esta terminal abierta mientras ejecutas la API.")

# Endpoints
doc.add_heading("Endpoints de la API", level=1)

doc.add_heading("1. GET /status - Health Check", level=2)
doc.add_paragraph("Verifica si la API está funcionando correctamente.")
p = doc.add_paragraph()
run = p.add_run("URL: http://localhost:5002/status")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("2. POST /chat - Consulta a la IA", level=2)
doc.add_paragraph("Envía un prompt a la IA y recibe una respuesta.")
p = doc.add_paragraph()
run = p.add_run("URL: http://localhost:5002/chat")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph("Body esperado:")
p = doc.add_paragraph()
run = p.add_run('{"prompt": "Tu pregunta aquí"}')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("3. GET / - Home", level=2)
doc.add_paragraph("Endpoint raíz para verificar que la API está respondiendo.")

# Ejecución
doc.add_heading("Ejecutar la API", level=1)

doc.add_heading("Terminal 1 - Iniciar Ollama:", level=2)
p = doc.add_paragraph()
run = p.add_run("ollama serve")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Terminal 2 - Iniciar API Flask:", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "cd C:\\xampp\\htdocs\\NOVA_ASISTENTE\\API_IA\n"
    ".\\venv\\Scripts\\activate\n"
    "python api_ia.py"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Logs
doc.add_heading("Logs y Monitoreo", level=1)

p = doc.add_paragraph()
run = p.add_run("Los logs se guardan en: logs/api_nova.log")
run.font.name = 'Courier New'

doc.add_heading("Ver últimas 50 líneas del log:", level=2)
p = doc.add_paragraph()
run = p.add_run("Get-Content logs\\api_nova.log -Tail 50")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Buscar errores:", level=2)
p = doc.add_paragraph()
run = p.add_run('Select-String "ERROR" logs\\api_nova.log')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Solución de Problemas
doc.add_heading("Solución de Problemas", level=1)

doc.add_heading("Problema 1: Connection refused", level=2)
doc.add_paragraph("Causa: Ollama no está ejecutándose")
p = doc.add_paragraph()
run = p.add_run("Solución: ")
run2 = p.add_run("ollama serve")
run2.font.name = 'Courier New'

doc.add_heading("Problema 2: Module not found", level=2)
doc.add_paragraph("Causa: El entorno virtual no está activado")
p = doc.add_paragraph()
run = p.add_run("Solución: ")
run2 = p.add_run(".\\venv\\Scripts\\activate")
run2.font.name = 'Courier New'

doc.add_heading("Problema 3: Puerto 5002 ya está en uso", level=2)
p = doc.add_paragraph()
run = p.add_run("netstat -ano | findstr :5002")
run.font.name = 'Courier New'
run.font.size = Pt(9)
p = doc.add_paragraph()
run = p.add_run("taskkill /PID [PID] /F")
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading("Problema 4: Modelo no encontrado", level=2)
p = doc.add_paragraph()
run = p.add_run("ollama pull qwen2.5:0.5b")
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Footer
doc.add_paragraph()
p = doc.add_paragraph(
    f"Documentacion generada: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')} | NOVA API IA v1.0"
)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.runs[0]
run.font.size = Pt(8)
run.font.italic = True

# Guardar documento
docx_file = "NOVA_API_IA_Documentacion.docx"
doc.save(docx_file)
print(f"[OK] DOCX generado exitosamente: {docx_file}")
print(f"[DOC] Ubicacion: C:\\xampp\\htdocs\\NOVA_ASISTENTE\\API_IA\\{docx_file}")
